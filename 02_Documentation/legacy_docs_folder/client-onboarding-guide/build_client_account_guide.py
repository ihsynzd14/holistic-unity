from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUTPUT = ROOT / "Holistic_Unity_Guida_Creazione_Account_Clienti_Web_iOS.docx"

BRAND = RGBColor(123, 34, 82)
TEXT = RGBColor(34, 34, 34)
MUTED = RGBColor(99, 99, 99)
LIGHT_FILL = "FBF6F8"
LABEL_FILL = "F4E8EE"

WEB_ONBOARDING_STEPS = [
    "Cosa ti porta qui, oggi?",
    "Cosa vorresti esplorare?",
    "Hai già esplorato qualcuna di queste pratiche?",
    "Quale approccio risuona di più?",
    "Quando senti di voler iniziare?",
    "In che fase ti senti?",
    "Cosa fa già parte della tua routine?",
    "C'è un segno che ti rappresenta? (opzionale)",
    "C'è qualcosa che vuoi farci sapere? (opzionale)",
]

IOS_ONBOARDING_STEPS = [
    "Welcome step: Inizia",
    "Intent",
    "Focus areas",
    "Familiar practices",
    "Approaches",
    "Timing",
    "Life season",
    "Current practices",
    "Cosmic marker (opzionale)",
    "Notes (opzionale)",
    "Ritual + recommendations + CTA finale",
]

WEB_STEP_CAPTIONS = [
    ("01-intent.png", "Web onboarding 01/10 - Intent"),
    ("02-focus-areas.png", "Web onboarding 02/10 - Focus areas"),
    ("03-familiar-practices.png", "Web onboarding 03/10 - Familiar practices"),
    ("04-approaches.png", "Web onboarding 04/10 - Approaches"),
    ("05-timing.png", "Web onboarding 05/10 - Timing"),
    ("06-life-season.png", "Web onboarding 06/10 - Life season"),
    ("07-current-practices.png", "Web onboarding 07/10 - Current practices"),
    ("08-cosmic-marker.png", "Web onboarding 08/10 - Cosmic marker"),
    ("09-notes.png", "Web onboarding 09/10 - Notes"),
    ("10-summary.png", "Web onboarding 10/10 - Summary"),
]

IOS_STEP_CAPTIONS = [
    ("01-onboarding-welcome.jpg", "iOS onboarding 01/12 - Welcome"),
    ("02-intent.jpg", "iOS onboarding 02/12 - Intent"),
    ("03-focus-areas.jpg", "iOS onboarding 03/12 - Focus areas"),
    ("04-familiar-practices.jpg", "iOS onboarding 04/12 - Familiar practices"),
    ("05-approaches.jpg", "iOS onboarding 05/12 - Approaches"),
    ("06-timing.jpg", "iOS onboarding 06/12 - Timing"),
    ("07-life-season.jpg", "iOS onboarding 07/12 - Life season"),
    ("08-current-practices.jpg", "iOS onboarding 08/12 - Current practices"),
    ("09-cosmic-marker.jpg", "iOS onboarding 09/12 - Cosmic marker"),
    ("10-notes.jpg", "iOS onboarding 10/12 - Notes"),
    ("11-ritual.jpg", "iOS onboarding 11/12 - Ritual"),
    ("12-summary.jpg", "iOS onboarding 12/12 - Summary and recommendations"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_language(document, lang="it-IT"):
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        rpr = style.element.get_or_add_rPr()
        lang_el = rpr.find(qn("w:lang"))
        if lang_el is None:
            lang_el = OxmlElement("w:lang")
            rpr.append(lang_el)
        lang_el.set(qn("w:val"), lang)


def configure_page(document):
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.16

    title = document.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = BRAND

    for name, size in (("Heading 1", 16), ("Heading 2", 12.5), ("Heading 3", 11)):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BRAND if name != "Heading 3" else TEXT
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)


def add_title_block(document):
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Guida completa creazione account clienti")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Web app client-side e iPhone app")
    run.bold = True
    run.font.size = Pt(12)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Dalla registrazione iniziale fino all'ingresso nel percorso personalizzato e alla dashboard."
    )
    run.font.color.rgb = MUTED
    run.font.size = Pt(10)

    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Aggiornato", "18 maggio 2026"),
        ("Scope", "Registrazione, conferma email, onboarding completo, accesso finale"),
        ("Web source", "Client web app in esecuzione su localhost + dashboard live"),
        (
            "iOS source",
            "Backup 6 Aprile: /Users/marcello/Desktop/Holistic Unity/iOS App/untitled folder/Backup 6 Aprile",
        ),
    ]
    for row, values in zip(table.rows, rows):
        for idx, value in enumerate(values):
            row.cells[idx].text = value
            for par in row.cells[idx].paragraphs:
                for run in par.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
            if idx == 0:
                set_cell_shading(row.cells[idx], LABEL_FILL)
    document.add_paragraph("")


def add_intro_box(document):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    run = p.add_run("Come usare questa guida")
    run.bold = True
    run.font.color.rgb = BRAND
    p = cell.add_paragraph(
        "Questa versione segue i flussi reali dell'app. Puoi usarla come guida per il team supporto, "
        "come base per una PDF cliente o come checklist durante un onboarding assistito."
    )
    p.paragraph_format.space_after = Pt(0)
    document.add_paragraph("")


def add_screenshot(document, filename, caption, width):
    path = ASSETS / filename
    if not path.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=width)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_numbered_step(document, title, body):
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    p.add_run(f" - {body}")


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def add_step_table(document, heading, rows):
    document.add_heading(heading, level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Ordine"
    hdr[1].text = "Step"
    set_cell_shading(hdr[0], LABEL_FILL)
    set_cell_shading(hdr[1], LABEL_FILL)
    for idx, label in enumerate(rows, start=1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = label
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
    document.add_paragraph("")


def add_sequence_shots(document, folder_name, items, width):
    for filename, caption in items:
        add_screenshot(document, str(Path(folder_name) / filename), caption, width)


def build_web_section(document):
    document.add_heading("1. Web app client-side", level=1)
    document.add_paragraph(
        "Il percorso web corretto non finisce dopo il form. Il cliente passa da registrazione, "
        "schermata di conferma email, questionario iniziale su `/welcome` e infine dashboard."
    )

    add_screenshot(document, "web-register.png", "Registrazione web: schermata iniziale", Inches(3.2))
    add_numbered_step(
        document,
        "Apri la pagina di registrazione",
        "Invia il cliente alla pagina dedicata e chiedigli di usare un indirizzo email controllabile subito.",
    )
    add_numbered_step(
        document,
        "Compila i campi richiesti",
        "Nome completo, email, numero di telefono, password e conferma password.",
    )
    add_numbered_step(
        document,
        "Approva i consensi obbligatori",
        "Il cliente deve accettare Termini + Privacy e l'approvazione specifica delle clausole onerose.",
    )
    add_numbered_step(
        document,
        "Clicca Create account",
        "Se il form e valido, la piattaforma crea l'account e prepara il passaggio successivo.",
    )

    add_screenshot(document, "web-check-email.png", "Dopo il submit: schermata Check your email", Inches(3.2))
    add_numbered_step(
        document,
        "Apri la mail di conferma",
        "Il cliente vede la schermata Check your email e deve usare il link ricevuto per confermare l'indirizzo.",
    )
    add_numbered_step(
        document,
        "Rientra nell'app dal link di conferma",
        "Dopo il redirect, il cliente entra nel percorso `/welcome` e completa il questionario iniziale.",
    )

    add_step_table(document, "Questionario iniziale web (`/welcome`)", WEB_ONBOARDING_STEPS)
    add_sequence_shots(document, "web-steps", WEB_STEP_CAPTIONS, Inches(5.8))

    add_numbered_step(
        document,
        "Rivedi i suggerimenti finali",
        "Alla fine del questionario la piattaforma mostra pratiche consigliate, operatori suggeriti e consenso research opzionale.",
    )
    add_numbered_step(
        document,
        "Entra nella dashboard",
        "Dopo il completamento, il cliente arriva in dashboard e puo iniziare a esplorare professionisti e sessioni.",
    )

    add_screenshot(document, "web-dashboard.png", "Dashboard cliente dopo il primo accesso", Inches(5.8))

    document.add_heading("Note utili per il supporto", level=2)
    add_bullet(document, "Se il cliente non vede l'email, fagli controllare spam, promozioni e aggiornamenti.")
    add_bullet(document, "Se apre il link ma non entra, fagli riprovare nello stesso browser usato per la registrazione.")
    add_bullet(document, "Se il cliente si ferma a meta del questionario, il flow va ripreso e concluso prima di usare la dashboard.")


def build_ios_section(document):
    document.add_heading("2. iPhone app", level=1)
    document.add_paragraph(
        "Questa sezione usa il backup iOS del 6 aprile che mi hai indicato. Il cliente parte dalla welcome screen, "
        "entra in Create Account, completa il questionario e vede un finale con raccomandazioni personalizzate."
    )

    add_screenshot(document, "ios-backup-welcome.jpg", "iPhone backup: welcome screen", Inches(2.45))
    add_numbered_step(
        document,
        "Apri l'app e tocca Get Started",
        "La schermata iniziale presenta il brand, i benefici principali e il punto d'ingresso verso la registrazione.",
    )

    add_screenshot(document, "ios-backup-create-account.jpg", "iPhone backup: schermata Create Account", Inches(2.45))
    add_numbered_step(
        document,
        "Scegli il metodo di accesso",
        "Il cliente puo continuare con Apple, Google oppure creare l'account con nome, email e password.",
    )
    add_numbered_step(
        document,
        "Completa la creazione account",
        "Dopo il sign up l'app porta il cliente dentro il questionario iniziale client-side.",
    )

    add_screenshot(document, "ios-backup-onboarding-welcome.jpg", "Welcome step dell'onboarding iOS", Inches(2.45))
    add_numbered_step(
        document,
        "Avvia il questionario",
        "La prima schermata dedicata all'onboarding introduce il percorso e anticipa che richiede circa 90 secondi.",
    )

    add_screenshot(document, "ios-backup-intent-step.jpg", "Step 1/9 dell'onboarding iOS: Intent", Inches(2.45))
    add_numbered_step(
        document,
        "Procedi step by step",
        "L'app mostra una progressione 1/9, 2/9 e cosi via fino all'ultimo passaggio opzionale prima del rituale finale.",
    )

    add_step_table(document, "Sequenza onboarding iOS (backup 6 Aprile)", IOS_ONBOARDING_STEPS)
    add_sequence_shots(document, "ios-steps", IOS_STEP_CAPTIONS, Inches(2.55))

    add_screenshot(document, "ios-backup-ritual.jpg", "Ritual step finale prima delle raccomandazioni", Inches(2.45))
    add_numbered_step(
        document,
        "Concludi con rituale e suggerimenti",
        "Alla fine l'utente passa in una schermata emozionale, poi vede pratiche consigliate, operatori suggeriti e CTA finale per entrare nel proprio percorso.",
    )
    add_numbered_step(
        document,
        "Accetta o salta le notifiche",
        "L'app puo chiedere il permesso notifiche nel tratto finale; il cliente puo consentire oppure rimandare.",
    )
    add_numbered_step(
        document,
        "Entra nell'esperienza principale",
        "Dopo la CTA finale il cliente atterra nell'area principale dell'app e puo iniziare a esplorare.",
    )

    document.add_heading("Cose da spiegare al cliente", level=2)
    add_bullet(document, "Apple e Google velocizzano l'accesso, ma il questionario iniziale va comunque completato.")
    add_bullet(document, "Cosmic marker e note sono opzionali; il cliente puo saltarli senza bloccare il percorso.")
    add_bullet(document, "Il finale mostra raccomandazioni personalizzate: vale la pena fermarsi un momento e leggerle prima di proseguire.")


def build_final_section(document):
    document.add_heading("3. Script pronto da inviare", level=1)
    document.add_paragraph(
        "Puoi usare questo testo in chat, email o WhatsApp quando accompagni un nuovo cliente."
    )
    quote = document.add_paragraph()
    quote.style = "Intense Quote" if "Intense Quote" in [s.name for s in document.styles] else "Normal"
    quote.add_run(
        "Apri Holistic Unity dal web o dall'app, crea il tuo account con email oppure con Apple o Google, "
        "completa i consensi richiesti e segui tutte le schermate del questionario iniziale. "
        "Alla fine vedrai i primi suggerimenti personalizzati e potrai entrare nella dashboard per iniziare a esplorare i professionisti disponibili."
    )

    document.add_heading("Chiusura rapida", level=2)
    add_bullet(document, "Web: form -> check email -> conferma link -> `/welcome` -> dashboard.")
    add_bullet(document, "iOS: welcome -> create account -> onboarding completo -> rituale/suggerimenti -> area principale.")
    add_bullet(document, "Se vuoi trasformarla in PDF cliente, questa base e gia pronta per una revisione grafica leggera.")


def main():
    document = Document()
    set_doc_language(document)
    configure_page(document)
    configure_styles(document)
    add_title_block(document)
    add_intro_box(document)
    build_web_section(document)
    document.add_section(WD_SECTION_START.NEW_PAGE)
    build_ios_section(document)
    build_final_section(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
